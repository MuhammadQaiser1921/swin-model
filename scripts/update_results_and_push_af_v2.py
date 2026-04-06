import argparse
import os
import subprocess
import sys
from datetime import datetime


def run(cmd, cwd):
    print("$", " ".join(cmd))
    subprocess.run(cmd, cwd=cwd, check=True)


def parse_args():
    parser = argparse.ArgumentParser(
        description="Update Research_Results.docx with AF_V2 audio metrics and push AF_V2 only."
    )
    parser.add_argument("--repo", default=os.path.abspath(os.path.join(os.path.dirname(__file__), "..")))
    parser.add_argument("--train-acc", type=float, required=True)
    parser.add_argument("--train-loss", type=float, required=True)
    parser.add_argument("--val-acc", type=float, required=True)
    parser.add_argument("--val-loss", type=float, required=True)
    parser.add_argument("--test-acc", type=float, required=True)
    parser.add_argument("--test-loss", type=float, required=True)
    parser.add_argument("--note", default="Audio model trained on Kaggle T4")
    parser.add_argument("--commit-message", default="AF_V2: add audio training results to Research_Results.docx")
    return parser.parse_args()


def ensure_docx_installed():
    try:
        import docx  # noqa: F401
    except Exception:
        print("python-docx not found. Installing...")
        subprocess.run([sys.executable, "-m", "pip", "install", "python-docx"], check=True)


def ensure_branch(repo_path):
    branch = subprocess.check_output(
        ["git", "branch", "--show-current"], cwd=repo_path, text=True
    ).strip()
    if branch != "AF_V2":
        raise RuntimeError(f"Current branch is '{branch}'. Switch to AF_V2 first.")


def update_docx(docx_path, metrics, note):
    from docx import Document

    doc = Document(docx_path)
    doc.add_paragraph("")
    doc.add_heading("AF_V2 - Audio Model Results", level=2)

    doc.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"Branch: AF_V2")
    doc.add_paragraph("Activations: MLP=Swish, Attention=Softmax, Output=Sigmoid")
    doc.add_paragraph(f"Note: {note}")

    table = doc.add_table(rows=7, cols=2)
    table.style = "Table Grid"

    rows = [
        ("Metric", "Value"),
        ("Train Accuracy", str(metrics['train_acc'])),
        ("Train Loss", str(metrics['train_loss'])),
        ("Validation Accuracy", str(metrics['val_acc'])),
        ("Validation Loss", str(metrics['val_loss'])),
        ("Test Accuracy", str(metrics['test_acc'])),
        ("Test Loss", str(metrics['test_loss'])),
    ]

    for i, (k, v) in enumerate(rows):
        table.cell(i, 0).text = k
        table.cell(i, 1).text = v

    doc.save(docx_path)


def main():
    args = parse_args()
    repo = os.path.abspath(args.repo)
    docx_path = os.path.join(repo, "Research_Results.docx")

    if not os.path.exists(docx_path):
        raise FileNotFoundError(f"Could not find {docx_path}")

    ensure_docx_installed()
    ensure_branch(repo)

    metrics = {
        "train_acc": args.train_acc,
        "train_loss": args.train_loss,
        "val_acc": args.val_acc,
        "val_loss": args.val_loss,
        "test_acc": args.test_acc,
        "test_loss": args.test_loss,
    }

    update_docx(docx_path, metrics, args.note)

    run(["git", "add", "Research_Results.docx"], cwd=repo)

    diff_cached = subprocess.check_output(["git", "diff", "--cached", "--name-only"], cwd=repo, text=True).strip()
    if not diff_cached:
        print("No staged changes found. Nothing to commit.")
        return

    run(["git", "commit", "-m", args.commit_message], cwd=repo)
    run(["git", "push", "origin", "AF_V2"], cwd=repo)

    print("Done: Research_Results.docx updated and pushed to AF_V2.")


if __name__ == "__main__":
    main()
